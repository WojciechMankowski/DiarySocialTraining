from docx import Document

class DOC:
    def __init__(self):
        self.document = Document()
        self.document.add_heading('Dziennik aktywności', 0)
    def addParagraph(self):
        p = self.document.add_paragraph('Poniżej znajdziesz Twoje wpisy: ')
    def AddTabet(self,  records):
        table = self.document.add_table(rows=1, cols=3)
        for qty, id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc

    def Save(self):
        self.document.save('wpis/Dziennik.docx')
